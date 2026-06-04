"""
word_tables_to_excel.py
-----------------------
Extrae tablas de datos de .docx y las exporta a Excel conservando:
  - Fuente, tamaño, negrita, itálica y color de texto
  - Fill (color de fondo) de celdas
  - Merges horizontales y verticales

Rutas de entrada/salida se guardan en config.json (misma carpeta que el script).
Primera vez: el programa las pide. Para cambiarlas después:
    python word_tables_to_excel.py --cambiar-rutas

Uso normal:
    python word_tables_to_excel.py
"""

import sys
import re
import json
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ── Configuración persistente ──────────────────────────────────────────────────

CONFIG_FILE = Path(__file__).parent / "config.json"


def load_config() -> dict:
    if CONFIG_FILE.exists():
        try:
            return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def save_config(cfg: dict):
    CONFIG_FILE.write_text(json.dumps(cfg, indent=2, ensure_ascii=False), encoding="utf-8")


def ask_path(prompt: str, must_exist: bool = True) -> Path:
    while True:
        raw = input(prompt).strip().strip('"').strip("'")
        p = Path(raw)
        if must_exist:
            if not p.exists():
                print("  ⚠  No existe esa ruta, intenta de nuevo.")
                continue
            if not p.is_dir():
                print("  ⚠  Esa ruta no es una carpeta.")
                continue
        return p


def get_paths(force_ask: bool = False) -> tuple[Path, Path]:
    cfg = load_config()

    if not force_ask and cfg.get("input_dir") and cfg.get("output_dir"):
        input_dir  = Path(cfg["input_dir"])
        output_dir = Path(cfg["output_dir"])
        if input_dir.exists():
            print(f"📁  Entrada : {input_dir}")
            print(f"📂  Salida  : {output_dir}")
            return input_dir, output_dir
        else:
            print("⚠  La carpeta de entrada guardada ya no existe. Se pedirán de nuevo.")

    print("\nConfigura las rutas (se guardarán para futuras ejecuciones):\n")
    input_dir  = ask_path("📁  Carpeta de entrada (donde están los .docx):\n> ", must_exist=True)
    output_dir = ask_path("📂  Carpeta de salida  (donde guardar los .xlsx):\n> ", must_exist=False)
    output_dir.mkdir(parents=True, exist_ok=True)

    save_config({"input_dir": str(input_dir), "output_dir": str(output_dir)})
    print("\n✅  Rutas guardadas en config.json\n")
    return input_dir, output_dir


# ── Constantes ─────────────────────────────────────────────────────────────────

EMU_PER_PT   = 12700
FALLBACK_FONT = "Garamond"
BORDER_SIDE  = Side(style="thin", color="BFBFBF")
CELL_BORDER  = Border(
    left=BORDER_SIDE, right=BORDER_SIDE,
    top=BORDER_SIDE,  bottom=BORDER_SIDE,
)


# ── Helpers de color ───────────────────────────────────────────────────────────

def hex_to_aRGB(hex_str: str) -> str:
    """Normaliza a AARRGGBB para openpyxl."""
    h = (hex_str or "").upper().lstrip("#")
    if len(h) == 6:
        return "FF" + h
    if len(h) == 8:
        return h
    return "FF000000"


def is_white_or_empty(aRGB: str) -> bool:
    rgb = aRGB[-6:].upper()
    return rgb in ("FFFFFF", "000000") and aRGB[:2] == "00"  # transparente


# ── Lectura de celdas Word ────────────────────────────────────────────────────

def cell_shading_hex(cell) -> str | None:
    """Fill de la celda (hex 6 chars) o None si no tiene / es blanco."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = (shd.get(qn("w:fill")) or "").upper()
    if fill in ("", "AUTO", "FFFFFF", "NONE"):
        return None
    return fill


def run_color_hex(run) -> str | None:
    """
    Color del run en hex 6 chars, o None si es heredado / automático.
    Devuelve None también para blanco puro (FFFFFF) en celdas sin fill oscuro,
    para evitar texto invisible.
    """
    try:
        if run.font.color and run.font.color.type is not None:
            rgb = str(run.font.color.rgb).upper()
            if rgb not in ("FFFFFF", "AUTO", "000000"):
                return rgb
            if rgb == "FFFFFF":
                return "FFFFFF"   # se manejará al escribir
    except Exception:
        pass
    return None


def cell_dominant_run(cell):
    """
    Devuelve (texto_completo, props_dict).
    props_dict: bold, italic, size_pt, color_hex|None, font_name, fill_hex|None
    """
    full_text = " ".join(p.text for p in cell.paragraphs).strip()
    fill_hex  = cell_shading_hex(cell)

    best_run, best_len = None, 0
    for para in cell.paragraphs:
        for run in para.runs:
            if len(run.text) > best_len:
                best_len = len(run.text)
                best_run = run

    if best_run is None:
        return full_text, {
            "bold": False, "italic": False, "size": 10,
            "color": None, "name": FALLBACK_FONT, "fill": fill_hex,
        }

    size_pt = (best_run.font.size / EMU_PER_PT) if best_run.font.size else 10
    color   = run_color_hex(best_run)

    # Si el color del run es blanco pero la celda NO tiene fill oscuro → usar negro
    if color == "FFFFFF" and not fill_hex:
        color = None   # → negro por defecto al escribir

    return full_text, {
        "bold":   bool(best_run.bold),
        "italic": bool(best_run.italic),
        "size":   size_pt,
        "color":  color,
        "name":   best_run.font.name or FALLBACK_FONT,
        "fill":   fill_hex,
    }


# ── Detección de merges ────────────────────────────────────────────────────────

def get_cell_merge_info(cell) -> tuple[int, str | None]:
    """
    Devuelve (gridSpan, vMerge_type).
    gridSpan: número de columnas que abarca (1 = sin merge horizontal).
    vMerge_type: 'restart' | 'continue' | None
    """
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 1, None

    gs = tcPr.find(qn("w:gridSpan"))
    span = int(gs.get(qn("w:val"), 1)) if gs is not None else 1

    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        vmerge = None
    else:
        vmerge = vm.get(qn("w:val"), "continue")

    return span, vmerge


def build_merge_map(table):
    """
    Construye una matriz de merge info para cada celda lógica.
    Devuelve list[list[dict]] con:
      - text, props  : contenido y formato
      - col_span     : número de cols que abarca (≥1)
      - row_span     : número de filas que abarca (≥1)
      - skip         : True si es celda continuación (no escribir)
      - xl_col       : columna Excel destino (1-based)
    
    python-docx repite celdas fusionadas en row.cells[], así que
    leemos directamente desde el XML para obtener las celdas únicas.
    """
    n_rows = len(table.rows)
    # grid_map[r][c] = None (libre) o (xl_r, xl_c) del merge que la ocupa
    grid_map = [[None] * 500 for _ in range(n_rows + 10)]

    result = []
    for r_idx, row in enumerate(table.rows):
        row_data = []
        xl_col   = 1  # columna Excel actual

        # Iterar sobre los <w:tc> del XML directamente (sin duplicados)
        tc_elements = row._tr.findall(qn("w:tc"))

        for tc in tc_elements:

            # Leer merge info del XML
            tcPr = tc.find(qn("w:tcPr"))
            gs   = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
            col_span = int(gs.get(qn("w:val"), 1)) if gs is not None else 1

            vm = tcPr.find(qn("w:vMerge")) if tcPr is not None else None
            if vm is None:
                vmerge = None
            else:
                vmerge = vm.get(qn("w:val"), "continue")

            # Texto y props (crear celda temporal para reutilizar helper)
            # Construir texto directamente desde el XML
            texts = []
            for p in tc.findall(qn("w:p")):
                t = "".join(r.text or "" for r in p.iter(qn("w:t")))
                if t.strip():
                    texts.append(t.strip())
            full_text = " ".join(texts)

            # Props: leer el run más largo del tc
            best_run_el, best_len = None, 0
            for run_el in tc.iter(qn("w:r")):
                t_el = run_el.find(qn("w:t"))
                if t_el is not None and len(t_el.text or "") > best_len:
                    best_len    = len(t_el.text)
                    best_run_el = run_el

            bold, italic, size_pt, color, font_name = False, False, 10, None, FALLBACK_FONT
            if best_run_el is not None:
                rPr = best_run_el.find(qn("w:rPr"))
                if rPr is not None:
                    bold   = rPr.find(qn("w:b"))  is not None
                    italic = rPr.find(qn("w:i"))  is not None
                    sz     = rPr.find(qn("w:sz"))
                    if sz is not None:
                        size_pt = int(sz.get(qn("w:val"), 20)) / 2
                    color_el = rPr.find(qn("w:color"))
                    if color_el is not None:
                        c = (color_el.get(qn("w:val")) or "").upper()
                        # Solo guardar si es un color real distinto de negro/auto
                        if c not in ("", "AUTO", "000000"):
                            color = c
                        # FFFFFF sólo si la celda tiene fill oscuro (se decide al escribir)
                    font_el = rPr.find(qn("w:rFonts"))
                    if font_el is not None:
                        font_name = (font_el.get(qn("w:ascii")) or
                                     font_el.get(qn("w:hAnsi")) or FALLBACK_FONT)

            # Fill de celda
            fill_hex = None
            if tcPr is not None:
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    f = (shd.get(qn("w:fill")) or "").upper()
                    if f not in ("", "AUTO", "FFFFFF", "NONE"):
                        fill_hex = f

            # Si color es FFFFFF pero no hay fill de color → negro (texto heredado)
            # Si color es None (no declarado) → negro por defecto
            if color == "FFFFFF" and not fill_hex:
                color = None  # → se escribirá como negro

            skip = (vmerge == "continue")

            if skip:
                # Celda continuación de merge vertical — no escribir, pero sí avanzar
                xl_col += col_span
                continue

            if not skip:
                # Marcar grid con vMerge y hSpan
                row_span = 1
                if vmerge == "restart":
                    # Contar cuántas filas continúan debajo
                    for rr in range(r_idx + 1, n_rows):
                        next_row_tcs = table.rows[rr]._tr.findall(qn("w:tc"))
                        # Buscar el tc en la misma posición de grid
                        grid_pos = xl_col
                        found_continue = False
                        cur_col = 1
                        for ntc in next_row_tcs:
                            while grid_map[rr][cur_col] is not None:
                                cur_col += 1
                            ntcPr = ntc.find(qn("w:tcPr"))
                            nvm   = ntcPr.find(qn("w:vMerge")) if ntcPr is not None else None
                            if cur_col == grid_pos and nvm is not None:
                                nvm_val = nvm.get(qn("w:val"), "continue")
                                if nvm_val != "restart":
                                    found_continue = True
                            ngs  = ntcPr.find(qn("w:gridSpan")) if ntcPr is not None else None
                            ncol = int(ngs.get(qn("w:val"), 1)) if ngs is not None else 1
                            cur_col += ncol
                        if found_continue:
                            row_span += 1
                        else:
                            break

                # Marcar celdas ocupadas en grid_map
                for dr in range(row_span):
                    for dc in range(col_span):
                        grid_map[r_idx + dr][xl_col + dc] = (r_idx + 1, xl_col)

                row_data.append({
                    "text":     full_text,
                    "bold":     bold,
                    "italic":   italic,
                    "size":     size_pt,
                    "color":    color,
                    "name":     font_name,
                    "fill":     fill_hex,
                    "col_span": col_span,
                    "row_span": row_span,
                    "skip":     False,
                    "xl_col":   xl_col,
                })
                xl_col += col_span

        result.append(row_data)
    return result


# ── Detección de tabla real / encabezado ────────────────────────────────────────

def is_real_table(table, min_cols=2, min_fill=0.3) -> bool:
    if len(table.columns) < min_cols:
        return False
    total, filled = 0, 0
    for row in table.rows:
        for tc in row._tr.findall(qn("w:tc")):
            total += 1
            t = "".join(r.text or "" for r in tc.iter(qn("w:t")))
            if t.strip():
                filled += 1
    return total > 0 and (filled / total) >= min_fill


def row_is_header(row) -> bool:
    """True si la mayoría de celdas tienen fill de color o negrita."""
    tcs    = row._tr.findall(qn("w:tc"))
    scored = 0
    for tc in tcs:
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is not None:
            shd = tcPr.find(qn("w:shd"))
            if shd is not None:
                f = (shd.get(qn("w:fill")) or "").upper()
                if f not in ("", "AUTO", "FFFFFF", "NONE"):
                    scored += 1
                    continue
        for rPr in tc.iter(qn("w:rPr")):
            if rPr.find(qn("w:b")) is not None:
                scored += 1
                break
    return scored >= max(1, len(tcs) // 2)


def detect_header_rows(table) -> int:
    rows = table.rows
    if not rows or not row_is_header(rows[0]):
        return 0
    if len(rows) > 1 and row_is_header(rows[1]):
        return 2
    return 1


# ── Iterador de elementos del documento ────────────────────────────────────────

def get_doc_elements(doc):
    body      = doc.element.body
    para_map  = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag == "p" and child in para_map:
            yield "paragraph", para_map[child]
        elif tag == "tbl" and child in table_map:
            yield "table", table_map[child]


# ── Nombre de hoja ──────────────────────────────────────────────────────────────

def safe_sheet_name(name: str, used: set, max_len=31) -> str:
    cleaned = re.sub(r'[\\/*?:\[\]]', '', name).strip()[:max_len] or "Tabla"
    original, counter = cleaned, 2
    while cleaned in used:
        suffix = f" ({counter})"
        cleaned = original[:max_len - len(suffix)] + suffix
        counter += 1
    used.add(cleaned)
    return cleaned


# ── Escritura en Excel ──────────────────────────────────────────────────────────

def autofit_columns(ws, min_w=8, max_w=55):
    for col_cells in ws.columns:
        length = max(len(str(c.value or "")) for c in col_cells)
        ws.column_dimensions[get_column_letter(col_cells[0].column)].width = \
            min(max(length + 2, min_w), max_w)


def write_sheet(ws, merge_map, n_header_rows):
    for r_idx, row_data in enumerate(merge_map, start=1):
        is_header = r_idx <= n_header_rows
        for cell_info in row_data:
            xl_col = cell_info["xl_col"]
            xl_cell = ws.cell(row=r_idx, column=xl_col, value=cell_info["text"])

            # Color de fuente
            color_hex = cell_info["color"]
            fill_hex  = cell_info["fill"]
            if color_hex:
                font_color = hex_to_aRGB(color_hex)
            elif fill_hex:
                # Celda con fondo de color pero sin color de texto declarado → blanco
                font_color = "FFFFFFFF"
            else:
                # Sin fill y sin color → negro siempre (datos o encabezados tipo formulario)
                font_color = "FF000000"

            xl_cell.font = Font(
                name   = cell_info["name"],
                size   = cell_info["size"],
                bold   = cell_info["bold"],
                italic = cell_info["italic"],
                color  = font_color,
            )

            xl_cell.alignment = Alignment(wrap_text=True, vertical="center",
                                          horizontal="center" if (cell_info["col_span"] > 1
                                                                  or cell_info["row_span"] > 1)
                                          else "left")
            xl_cell.border = CELL_BORDER

            # Fill
            if cell_info["fill"]:
                xl_cell.fill = PatternFill("solid",
                                           start_color=hex_to_aRGB(cell_info["fill"]))

            # Merge
            cs, rs = cell_info["col_span"], cell_info["row_span"]
            if cs > 1 or rs > 1:
                end_row = r_idx + rs - 1
                end_col = xl_col + cs - 1
                ws.merge_cells(
                    start_row=r_idx, start_column=xl_col,
                    end_row=end_row,  end_column=end_col,
                )

    autofit_columns(ws)
    if n_header_rows >= 1 and len(merge_map) > n_header_rows:
        ws.freeze_panes = f"A{n_header_rows + 1}"


# ── Procesamiento de un archivo ─────────────────────────────────────────────────

def process_docx(docx_path: Path, output_dir: Path) -> tuple[bool, str]:
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        return False, f"No se pudo abrir: {e}"

    wb = Workbook()
    wb.remove(wb.active)
    used_names   = set()
    tables_found = 0
    last_title   = ""

    for kind, obj in get_doc_elements(doc):
        if kind == "paragraph":
            text = obj.text.strip()
            if text:
                last_title = text
        elif kind == "table":
            if not is_real_table(obj):
                continue

            n_headers = detect_header_rows(obj)
            merge_map = build_merge_map(obj)
            if not merge_map:
                continue

            sheet_name = safe_sheet_name(
                last_title or f"Tabla {tables_found + 1}", used_names
            )
            ws = wb.create_sheet(title=sheet_name)
            write_sheet(ws, merge_map, n_headers)
            tables_found += 1

    if tables_found == 0:
        return False, "No se encontraron tablas de datos."

    xlsx_path = output_dir / (docx_path.stem + ".xlsx")
    wb.save(str(xlsx_path))

    new_name = docx_path.stem + "_CONVERTIDO" + docx_path.suffix
    docx_path.rename(docx_path.parent / new_name)

    return True, f"{tables_found} tabla(s) → {xlsx_path.name}"


# ── Main ────────────────────────────────────────────────────────────────────────

def main():
    force_ask = "--cambiar-rutas" in sys.argv

    print("=" * 55)
    print("  Word Tables → Excel Converter")
    print("=" * 55 + "\n")

    input_dir, output_dir = get_paths(force_ask)
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_files = [
        f for f in input_dir.glob("*.docx")
        if not f.stem.endswith("_CONVERTIDO")
    ]

    if not docx_files:
        print("\n⚠  No se encontraron archivos .docx en esa carpeta.")
        sys.exit(0)

    print(f"\n🔍  Archivos encontrados: {len(docx_files)}\n")
    print("-" * 55)

    ok_count = err_count = 0
    for docx_path in docx_files:
        print(f"  ⚙  {docx_path.name}  ...", end=" ", flush=True)
        success, msg = process_docx(docx_path, output_dir)
        if success:
            print(f"✅  {msg}")
            ok_count += 1
        else:
            print(f"❌  {msg}")
            err_count += 1

    print("-" * 55)
    print(f"\n✔  Convertidos: {ok_count}   ✘  Errores: {err_count}")
    print(f"📂  Salida en: {output_dir}\n")


if __name__ == "__main__":
    main()
